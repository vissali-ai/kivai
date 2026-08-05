from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

OUT = Path(r"C:\Users\Windows 11\kivai\Inventario-de-Ferramentas-Kivai.docx")
BLUE = "2E74B5"; DARK = "1F4D78"; PALE = "E8EEF5"; LIGHT = "F4F7FA"; GRAY = "667085"; WHITE = "FFFFFF"

groups = {
"Imagens (14)": [
("Removedor de Fundo","removedor-de-fundo","Remove automaticamente o fundo de imagens; possui motor no navegador e API Python/FastAPI."),
("Compressor de Imagens","compressor-de-imagens","Reduz arquivos JPG, PNG e WebP com controle de qualidade."),
("Conversor de Imagens","conversor-de-imagens","Converte entre PNG, JPG e WebP."),
("Conversor HEIC","conversor-heic","Converte HEIC/HEIF para JPG."),
("Redimensionar Imagem","redimensionar-imagem","Altera largura e altura, com preservação de proporção."),
("Gerador de QR Code","gerador-de-qr-code","Cria QR Codes para links, texto, WhatsApp, e-mail, telefone e Wi-Fi."),
("Gerador de Favicon","gerador-de-favicon","Gera favicon.ico e ícones PNG em tamanhos comuns."),
("Imagem → Favicon","imagem-para-favicon","Transforma uma imagem em pacote de favicons para web e dispositivos."),
("Gerador de Mockups","gerador-de-mockups","Aplica artes a mockups de dispositivos, camiseta, cartão, outdoor e redes sociais."),
("Recortar Imagem","recortar-imagem","Recorta uma área selecionada da imagem."),
("Adicionar Marca d'Água","adicionar-marca-dagua","Insere marca d'água textual ou visual em imagens."),
("Espelhar e Girar Imagem","espelhar-e-girar-imagem","Espelha horizontal/verticalmente e rotaciona imagens."),
("Conversor SVG ↔ PNG","conversor-svg-png","Converte SVG em PNG e encapsula imagens raster em SVG."),
("Gerador de Placeholder (LQIP)","gerador-de-placeholder","Gera uma prévia leve e de baixa qualidade para carregamento progressivo."),],
"PDFs (6)": [
("PDF para Imagens","pdf-para-imagens","Renderiza as páginas de um PDF como imagens."),("Imagens para PDF","imagens-para-pdf","Monta um PDF a partir de uma ou mais imagens."),("Unir PDFs","unir-pdfs","Combina vários PDFs em um único arquivo."),("Dividir PDF","dividir-pdf","Separa páginas ou intervalos em arquivos independentes."),("Girar PDF","girar-pdf","Rotaciona páginas de um documento PDF."),("Compactar PDF","compactar-pdf","Reduz o tamanho do PDF por recomposição/otimização no navegador."),],
"Calculadoras (6)": [
("Calculadora de ROAS","calculadora-de-roas","Calcula retorno sobre investimento em mídia."),("Calculadora de ROI","calculadora-de-roi","Calcula retorno financeiro sobre o investimento."),("Calculadora de Markup","calculadora-de-markup","Calcula preço ou multiplicador de markup."),("Calculadora de Margem","calculadora-de-margem","Calcula margem e relação entre custo, preço e lucro."),("Calculadora de Desconto","calculadora-de-desconto","Calcula preço final e economia após desconto."),("Calculadora de Porcentagem","calculadora-de-porcentagem","Resolve variação, parcela e relações percentuais."),],
"Texto (1)": [("Contador de Palavras","contador-de-palavras","Conta palavras, caracteres e métricas relacionadas ao texto.")],
"Vídeo (10)": [
("Remover áudio","remover-audio-video","Gera uma versão silenciosa em WebM."),("Girar vídeo","girar-video","Rotaciona em 90°, 180° ou 270° e exporta em WebM."),("Espelhar vídeo","espelhar-video","Espelha horizontal ou verticalmente."),("Recortar vídeo (Crop)","recortar-video","Recorta a área visível do vídeo."),("Gerador de Thumbnail","capturar-frame-video","Captura um frame e exporta em PNG ou JPG."),("Alterar volume do vídeo","alterar-volume-video","Aumenta, reduz ou remove o áudio."),("Ajustar velocidade do vídeo","ajustar-velocidade-video","Acelera ou desacelera vídeo e áudio."),("Vídeo para áudio","video-para-audio","Extrai a trilha de áudio em WebM."),("Dividir vídeo","dividir-video","Divide o vídeo em duas partes no ponto escolhido."),("Redimensionar vídeo","redimensionar-video","Altera resolução com modos ajustar, preencher ou esticar."),],
"Social media e conteúdo (36)": [
("Quebra de Linha para Instagram","quebra-de-linha-instagram","Formata legendas com quebras prontas para copiar."),("Contador de Caracteres Instagram","contador-de-caracteres-instagram","Conta caracteres e acompanha o limite da legenda."),("Contador de Hashtags Instagram","contador-de-hashtags-instagram","Conta hashtags, palavras e caracteres."),
("Baixar Thumbnail do YouTube","baixar-thumbnail-youtube","Obtém a thumbnail pública de um vídeo."),("Extrair Thumbnail de Shorts","extrair-thumbnail-shorts","Localiza a thumbnail pública de um YouTube Short."),("Simulador de Bio Instagram","simulador-de-bio-instagram","Pré-visualiza uma bio antes de publicar."),("Gerador de Link na Bio Instagram","gerador-de-link-na-bio-instagram","Cria links para bio e campanhas."),("Gerador de Hashtags Instagram","gerador-de-hashtags-instagram","Monta hashtags por categoria e objetivo."),("Gerador de Hashtags TikTok","gerador-de-hashtags-tiktok","Gera hashtags com modelos de conteúdo."),("Gerador de Títulos TikTok","gerador-de-titulos-tiktok","Cria títulos curtos com modelos."),("Contador de Caracteres Facebook","contador-de-caracteres-facebook","Conta caracteres, palavras e linhas."),("Gerador de Descrição Facebook","gerador-de-descricao-facebook","Cria descrições a partir de modelos."),("Gerador de Headline LinkedIn","gerador-de-headline-linkedin","Monta headlines profissionais."),("Gerador de Resumo Profissional LinkedIn","gerador-de-resumo-profissional-linkedin","Estrutura o resumo do perfil."),("Gerador de Título Pinterest","gerador-de-titulo-pinterest","Gera títulos claros e pesquisáveis."),("Gerador de Descrição Pinterest","gerador-de-descricao-pinterest","Cria descrições para Pins."),("Gerador de Alt Text Pinterest","gerador-de-alt-text-pinterest","Escreve textos alternativos descritivos."),("Contador de Caracteres Twitter (X)","contador-de-caracteres-twitter","Acompanha o limite do post."),("Gerador de Mensagens WhatsApp","gerador-de-mensagens-whatsapp","Cria mensagens com modelos."),("Gerador de Convites WhatsApp","gerador-de-convites-whatsapp","Gera convites para eventos e grupos."),("Gerador de CTA WhatsApp","gerador-de-cta-whatsapp","Cria chamadas para ação direcionadas ao WhatsApp."),("Gerador de Prompts","gerador-de-prompts","Monta prompts a partir de objetivo e contexto."),("Gerador de Emojis","gerador-de-emojis","Sugere combinações por tema e intenção."),("Gerador de CTA","gerador-de-cta","Cria chamadas para ação por objetivo."),("Gerador de Títulos","gerador-de-titulos","Gera ideias usando modelos estáticos."),("Gerador de Descrições","gerador-de-descricoes","Cria descrições objetivas com modelos."),("Gerador de Resumos","gerador-de-resumos","Estrutura resumos curtos no navegador."),("Gerador de Palavras-chave","gerador-de-palavras-chave","Cria listas por tema e intenção."),("Banco de Datas Comemorativas","banco-de-datas-comemorativas","Apoia o planejamento de conteúdo por calendário."),("Conversor de Texto","conversor-de-texto","Transforma, limpa e formata textos."),("Contador de Tempo de Leitura","contador-de-tempo-de-leitura","Estima a duração da leitura."),("Contador de Tempo de Fala","contador-de-tempo-de-fala","Estima a duração da fala."),("Calculadora de Engajamento","calculadora-de-engajamento","Calcula a taxa de engajamento."),("Calculadora de CPM","calculadora-de-cpm","Calcula custo por mil impressões."),("Calculadora de ROI para Influenciadores","calculadora-de-roi-influenciadores","Estima retorno de campanhas com influenciadores."),("Calculadora de Alcance Estimado","calculadora-de-alcance-estimado","Estima alcance por seguidores e taxa."),]
}

def shade(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=tcPr.find(qn('w:shd'))
    if shd is None: shd=OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'),fill)

def margins(cell, top=80, start=120, bottom=80, end=120):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr(); node=tcPr.first_child_found_in('w:tcMar')
    if node is None: node=OxmlElement('w:tcMar'); tcPr.append(node)
    for side,val in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        el=node.find(qn('w:'+side))
        if el is None: el=OxmlElement('w:'+side); node.append(el)
        el.set(qn('w:w'),str(val)); el.set(qn('w:type'),'dxa')

def set_width(cell, dxa):
    tcPr=cell._tc.get_or_add_tcPr(); tcW=tcPr.find(qn('w:tcW'))
    if tcW is None: tcW=OxmlElement('w:tcW'); tcPr.append(tcW)
    tcW.set(qn('w:w'),str(dxa)); tcW.set(qn('w:type'),'dxa')

def table_geometry(table, widths):
    table.autofit=False; tblPr=table._tbl.tblPr
    tblW=tblPr.find(qn('w:tblW'))
    if tblW is None: tblW=OxmlElement('w:tblW'); tblPr.append(tblW)
    tblW.set(qn('w:w'),str(sum(widths))); tblW.set(qn('w:type'),'dxa')
    tblInd=tblPr.find(qn('w:tblInd'))
    if tblInd is None: tblInd=OxmlElement('w:tblInd'); tblPr.append(tblInd)
    tblInd.set(qn('w:w'),'120'); tblInd.set(qn('w:type'),'dxa')
    grid=table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for w in widths:
        col=OxmlElement('w:gridCol'); col.set(qn('w:w'),str(w)); grid.append(col)
    for row in table.rows:
        for i,cell in enumerate(row.cells): set_width(cell,widths[i]); margins(cell); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER

def add_field(paragraph, instruction):
    run=paragraph.add_run(); fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),instruction); run._r.addnext(fld)

doc=Document(); sec=doc.sections[0]
sec.page_width=Inches(8.5); sec.page_height=Inches(11); sec.top_margin=sec.bottom_margin=sec.left_margin=sec.right_margin=Inches(1)
styles=doc.styles
normal=styles['Normal']; normal.font.name='Calibri'; normal.font.size=Pt(11); normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.25
for name,size,color,before,after in [('Title',28,DARK,0,8),('Subtitle',12,GRAY,0,18),('Heading 1',16,BLUE,18,10),('Heading 2',13,BLUE,14,7),('Heading 3',12,DARK,10,5)]:
    st=styles[name]; st.font.name='Calibri'; st.font.size=Pt(size); st.font.color.rgb=RGBColor.from_string(color); st.font.bold=name!='Subtitle'; st.paragraph_format.space_before=Pt(before); st.paragraph_format.space_after=Pt(after); st.paragraph_format.keep_with_next=True
header=sec.header.paragraphs[0]; header.text='KIVAI  |  INVENTÁRIO TÉCNICO'; header.style=styles['Caption']; header.runs[0].font.color.rgb=RGBColor.from_string(GRAY)
footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER; footer.add_run('Kivai • Inventário de ferramentas • Página '); add_field(footer,'PAGE')

p=doc.add_paragraph(style='Title'); p.add_run('Inventário de Ferramentas do Kivai')
p=doc.add_paragraph(style='Subtitle'); p.add_run('Catálogo funcional e tecnológico do projeto • 4 de agosto de 2026')
p=doc.add_paragraph(); p.add_run('Escopo. ').bold=True; p.add_run('Levantamento feito diretamente no registro central, nas rotas e nos motores do repositório. O catálogo separa as ferramentas oferecidas ao usuário da infraestrutura usada para construí-las.')
t=doc.add_table(rows=2, cols=6); t.alignment=WD_TABLE_ALIGNMENT.CENTER
for i,(n,v) in enumerate([('Total','73'),('Imagens','14'),('PDF','6'),('Calculadoras','6'),('Social','36'),('Vídeo + texto','11')]):
    c=t.cell(0,i); c.text=v; shade(c,BLUE); c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER; c.paragraphs[0].runs[0].font.bold=True; c.paragraphs[0].runs[0].font.color.rgb=RGBColor.from_string(WHITE)
    c=t.cell(1,i); c.text=n; shade(c,PALE); c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
table_geometry(t,[1560]*6)

doc.add_heading('Leitura rápida',level=1)
for text in [
'Todas as 73 ferramentas estão marcadas como disponíveis no registro lib/tools.ts.',
'As rotas dedicadas cobrem imagem, PDF, calculadoras, texto, vídeo e três utilitários específicos de Instagram.',
'Outras 33 ferramentas sociais são atendidas pela rota dinâmica /ferramentas/[slug], usando um cliente compartilhado orientado pelo slug.',
'A maior parte do processamento de arquivos ocorre localmente no navegador; a remoção de fundo também possui backend Python dedicado.'
]: doc.add_paragraph(text,style='List Bullet')

doc.add_heading('Catálogo completo',level=1)
for heading,items in groups.items():
    doc.add_heading(heading,level=2)
    table=doc.add_table(rows=1,cols=3); table.style='Table Grid'; table.rows[0]._tr.get_or_add_trPr().append(OxmlElement('w:tblHeader'))
    for i,h in enumerate(['Ferramenta','Slug / rota','Função principal']):
        cell=table.rows[0].cells[i]; cell.text=h; shade(cell,PALE); cell.paragraphs[0].runs[0].font.bold=True
    for name,slug,desc in items:
        cells=table.add_row().cells; cells[0].text=name; cells[1].text='/ferramentas/'+slug; cells[2].text=desc
        for c in cells:
            for p in c.paragraphs:
                p.paragraph_format.space_after=Pt(2); p.paragraph_format.line_spacing=1.08
                for r in p.runs: r.font.size=Pt(9.2)
    table_geometry(table,[2500,2700,4160])

doc.add_heading('Tecnologias e bibliotecas do projeto',level=1)
tech=[
('Aplicação web','Next.js 16.2.10, React 19.2.4 e TypeScript 5.','Estrutura App Router, páginas, rotas dinâmicas e componentes.'),
('Interface','Tailwind CSS 4, shadcn, Radix UI, Framer Motion, Lucide e Phosphor Icons.','Layout, componentes, animação e iconografia.'),
('Imagens','Canvas/Web APIs, heic2any, JSZip e motores próprios em lib/.','Conversão, compressão, corte, transformação, favicon e mockups.'),
('PDF','pdf-lib, pdfjs-dist e jsPDF.','Leitura, renderização, montagem, união, divisão, giro e compactação.'),
('QR Code','qrcode.','Geração de QR Codes no cliente.'),
('Remoção de fundo','FastAPI, Uvicorn, Pillow, python-multipart, PyTorch, Transformers, rembg e ONNX Runtime.','API Python e inferência de segmentação; há implementações RMBG/BiRefNet no frontend.'),
('Publicidade e consentimento','Google AdSense/Next Third Parties e módulos próprios de consentimento/cookies.','Monetização, privacidade e preferências.'),
('Qualidade e build','ESLint 9, eslint-config-next, PostCSS e scripts npm dev/build/start/lint.','Desenvolvimento, compilação e análise estática.'),]
table=doc.add_table(rows=1,cols=3); table.style='Table Grid'
for i,h in enumerate(['Área','Ferramentas','Uso no Kivai']): table.rows[0].cells[i].text=h; shade(table.rows[0].cells[i],PALE); table.rows[0].cells[i].paragraphs[0].runs[0].font.bold=True
for row in tech:
    cells=table.add_row().cells
    for i,val in enumerate(row): cells[i].text=val
table_geometry(table,[1800,3600,3960])

doc.add_heading('Arquitetura funcional observada',level=1)
for title,body in [
('Registro central','lib/tools.ts define categorias, metadados, disponibilidade, SEO e slugs.'),
('Rotas dedicadas','app/ferramentas/<slug>/page.tsx e clientes associados implementam fluxos específicos.'),
('Rota dinâmica','app/ferramentas/[slug]/page.tsx atende o catálogo social baseado em definições e templates compartilhados.'),
('Motores reutilizáveis','lib/ concentra lógica para anúncios, imagens, favicon, QR Code, mockups, conversão e remoção de fundo.'),
('Backend especializado','backend/main.py expõe o processamento de remoção de fundo e dependências de inferência.'),
]:
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(5); r=p.add_run(title+': '); r.bold=True; r.font.color.rgb=RGBColor.from_string(DARK); p.add_run(body)

doc.add_heading('Fontes do inventário',level=1)
for f in ['lib/tools.ts — catálogo e status das ferramentas.','app/ferramentas/ — páginas dedicadas, hubs e rota dinâmica.','components/tools/ e lib/ — clientes, shells e motores reutilizáveis.','package.json — dependências JavaScript/TypeScript.','backend/main.py e backend/requirements.txt — API e dependências Python.']:
    doc.add_paragraph(f,style='List Bullet')
p=doc.add_paragraph(); p.add_run('Nota de escopo: ').bold=True; p.add_run('“Ferramenta” neste documento significa funcionalidade oferecida ao usuário final. Páginas institucionais, componentes de UI, hubs, serviços comerciais e arquivos de teste não foram contados como ferramentas.')

doc.core_properties.title='Inventário de Ferramentas do Kivai'; doc.core_properties.subject='Catálogo funcional e tecnológico'; doc.core_properties.author='Kivai'
doc.save(OUT)
print(OUT)
